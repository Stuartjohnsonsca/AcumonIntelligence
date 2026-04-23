"""
Convert the generated markdown test-spec files into a single human-readable
Word document.

Output: test-specs.docx (section per FS line, proper headings, TOC-ready)
"""
from __future__ import annotations
import glob
import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "test-specs.docx")

FS_LINE_ORDER = [
    "accruals", "amount-owed-by-group-undertakings", "amounts-owed-to-group-undertakings",
    "cash-and-cash-equivalents", "corporation-tax-payable", "cost-of-sales",
    "deferred-revenue", "deferred-tax", "going-concern", "intangible-assets",
    "interest-payable-and-similar-income", "inventory", "investments-financial-assets",
    "investments-in-subsidaries", "loans-and-borrowings", "management-override",
    "notes-and-disclosures", "operating-expenses", "other-creditors", "other-debtors",
    "other-interest-receivable-and-similar-income", "other-operating-income",
    "other-taxation-and-social-security-payable", "prepayments-and-accrued-income",
    "property-plant-and-equipment", "reserves", "revenue", "share-capital",
    "tax-expense", "trade-creditors", "trade-debtors", "wages-salaries",
]

# ─── Markdown → docx helpers ──────────────────────────────────────────────

INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ICON_RE = re.compile(r"(🆕)")

def add_runs(paragraph, text: str):
    """
    Split a line on inline markdown (bold, inline code) and emit runs with
    appropriate formatting. Leaves other characters as plain text.
    """
    # Token stream: find all matches of bold/code/emoji, split around them.
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|🆕)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        elif tok == "🆕":
            r = paragraph.add_run(" [NEW] ")
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); r.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def strip_md_inline(s: str) -> str:
    s = BOLD_RE.sub(r"\1", s)
    s = INLINE_CODE_RE.sub(r"\1", s)
    return s

def render_md_file(doc: Document, md_path: str):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    in_blockquote = False
    for raw in lines:
        line = raw.rstrip()

        if not line.strip():
            doc.add_paragraph("")
            in_blockquote = False
            continue

        # Horizontal rule
        if line.strip() in ("---", "***"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # Headings
        if line.startswith("### "):
            p = doc.add_heading(level=3)
            add_runs(p, line[4:])
            continue
        if line.startswith("## "):
            p = doc.add_heading(level=2)
            add_runs(p, line[3:])
            continue
        if line.startswith("# "):
            p = doc.add_heading(level=1)
            add_runs(p, line[2:])
            continue

        # Blockquote (used for "Full description")
        if line.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, line[2:])
            in_blockquote = True
            continue

        # Numbered list (" 1. ...")
        m_num = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m_num.group(2))
            continue

        # Bulleted list (" - ..." or "- ...")
        m_bul = re.match(r"^\s*-\s+(.*)$", line)
        if m_bul:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m_bul.group(1))
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_runs(p, line)

def set_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

def add_cover(doc: Document):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Audit Test Specs — Draft Test Pack")
    r.bold = True
    r.font.size = Pt(24)
    doc.add_paragraph("")
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("534 tests across 32 FS lines · FRS 102 · Significant Risk")
    rs.italic = True
    rs.font.size = Pt(12)
    rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph("")
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = d.add_run("Each spec block lists the Action Pipeline steps for Claude to build. [NEW] marks Actions not yet in production.")
    rd.font.size = Pt(10)
    rd.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.add_page_break()

def main():
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    add_cover(doc)

    # Render README first
    render_md_file(doc, os.path.join(HERE, "README.md"))
    doc.add_page_break()

    # New Actions catalogue
    render_md_file(doc, os.path.join(HERE, "_new-actions.md"))
    doc.add_page_break()

    # Per FS line — in the defined order, page-break between each
    for i, slug in enumerate(FS_LINE_ORDER):
        md = os.path.join(HERE, f"{slug}.md")
        if not os.path.exists(md):
            print(f"  [missing] {slug}.md")
            continue
        render_md_file(doc, md)
        if i < len(FS_LINE_ORDER) - 1:
            doc.add_page_break()

    doc.save(OUT_PATH)
    size_kb = os.path.getsize(OUT_PATH) // 1024
    print(f"Wrote {OUT_PATH} ({size_kb} KB)")

if __name__ == "__main__":
    main()
